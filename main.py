import scrapper
import handle_embeddings
import configparser
import openai
from docx import Document

config = configparser.ConfigParser()
config.read('config.ini')

CSV_FILE_NAME = config.get('SCRAPPER', 'CSV_FILE_NAME')

def generate_report(values, query):
    api_key = config.get('SECRETS', 'OPENAI_API_KEY')
    client = openai.OpenAI(api_key=api_key)
    prompt_template = """
    Articles:
    ###
    First Article:
    {}

    Second Article:
    {}

    Third Article:
    {}

    ###

    Answer the following question using only information from the articles. 
    Answer in a complete sentence, with proper capitalization and punctuation. 
    Answer in 500 to 2000 words, provide detailed imformation, elaborate properly.
    If there is no good answer in the article, say "I don't know".

    Question: {}
    Answer: 
    """
    prompt_template = prompt_template.format(values[0], values[1], values[2], query)

    try:
        response = client.completions.create(
            model='gpt-3.5-turbo',
            prompt=prompt_template,
            # n=1
        )
        return response['choices']['text']
    except openai.RateLimitError as rate_limit_error:
        print('Upgrade your OpenAI plan to use this service, current plan hits rate limit')
    except Exception as e:
        print(f'Error occured while generating report\nError: \n{e}')


def save_report(data):
    doc = Document()
    doc.add_heading('Report', level=1)
    for item in data:
        doc.add_heading('Question:', level=2)
        doc.add_paragraph(item['query'])
        doc.add_heading('Answer:', level=2)
        doc.add_paragraph(item['answer'])
        doc.add_paragraph('\n')

    doc.save('report.docx')

    print('Successfully generated report filename - "report.docx"')

if __name__ == "__main__":
    mode = config.get('MODE', 'MODE')
    queries = [
        "Identify the industry in which Canoo operates, along with its size, growth rate, trends, and key players. ",
        "Analyze Canoo's main competitors, including their market share, products or services offered, pricing strategies, and marketing efforts. ",
        "Identify key trends in the market, including changes in consumer behavior, technological advancements, and shifts in the competitive landscape. ",
        "Gather information on Canoo's financial performance, including its revenue, profit margins, return on investment, and expense structure."
    ]
    if mode == "SEARCH_AND_BUILD":
        scrapper.start()
        handle_embeddings.generate_embedding_knowledgebase(CSV_FILE_NAME)
    report_data = []
    for query in queries:
        result = handle_embeddings.perform_vector_search(query)
        values = [d['content'] for d in result]
        ans = generate_report(values, query)
        report_data.append({'query': query, 'answer': ans})

    save_report(report_data)
